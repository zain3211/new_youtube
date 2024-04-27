import streamlit as st
import os
from pytube import YouTube
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import StringIO
from reportlab.lib.pagesizes import letter
import docx
from youtube_transcript_api import YouTubeTranscriptApi
from datetime import timedelta
from docx.shared import Pt 
from youtube_transcript_api import YouTubeTranscriptApi
import requests
import time 
import assemblyai as aai
from pytube import YouTube
import assemblyai


def extract_video_id(video_url):
    try:
        # Extract video ID from the YouTube URL
        video_id = YouTube(video_url).video_id
        return video_id
    except Exception as e:
        st.error(f"Error extracting video ID: {str(e)}")
        return None

# Download YouTube video
def download_youtube_video(video_url, selected_quality):
    try:
        yt = YouTube(video_url)
        available_streams = yt.streams.filter(res=selected_quality)

        if not available_streams:
            st.error(f"No streams available for the selected quality ({selected_quality}). Try another quality.")
            return None

        best_stream = available_streams.first()  # Pick the first available stream
        download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        file_path = best_stream.download(download_dir)

        st.success(f"Video downloaded successfully at: {file_path}")
        return file_path
    except Exception as e:
        st.error(f"Error downloading video: {str(e)}")
        return None
# Download YouTube audio
def download_youtube_audio(video_url):
    try:
        yt = YouTube(video_url)
        audio_stream = yt.streams.filter(only_audio=True).first()
        download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        sanitized_title = re.sub(r'[<>:"/\\|?*]', '', yt.title)
        audio_file_path = os.path.join(download_dir, f"{sanitized_title}.mp3")
        audio_stream.download(download_dir, filename=f"{sanitized_title}.mp3")

        st.success(f"Audio downloaded successfully at: {audio_file_path}")
        return audio_file_path
    except Exception as e:
        st.error(f"Error downloading audio: {str(e)}")
        return None
def send_audio_to_assembly_ai(audio_file_path):
  # AssemblyAI API key
  aai.settings.api_key = '6d7723dc0ade4f38af10feaf48ef637e'

  try:
    # Create a transcriber object
    transcriber = assemblyai.Transcriber()

    # Transcribe the audio file
    transcript = transcriber.transcribe(audio_file_path)

    if transcript.status == assemblyai.TranscriptStatus.error:
      st.error(f"Transcription failed: {transcript.error}")
      return None
    else:
      return transcript.text
  except Exception as e:
    st.error(f"Error using AssemblyAI: {str(e)}")
    return None

def get_text_as_word_doc(video_url):
  try:
    video_id = extract_video_id(video_url)
    if not video_id:
      return None

    try:
      transcript = YouTubeTranscriptApi.get_transcript(video_id)
    except Exception:
      # If the transcript is not available through the API, download the audio and use AssemblyAI
      audio_path = download_youtube_audio(video_url)
      if not audio_path:
        raise Exception("Failed to download audio for transcription.")
      transcript_text = send_audio_to_assembly_ai(audio_path)
      if not transcript_text:
        raise Exception("Failed to get transcript from audio.")
      transcript = [{'text': transcript_text, 'start': 0, 'duration': 0}]  # Mock segment for compatibility

    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)

    # Add timestamp and text for each transcript segment
    for segment in transcript:
      if 'start' in segment and 'duration' in segment:
        start_time = segment['start']
        end_time = segment['start'] + segment['duration']
        start_time_str = str(timedelta(seconds=start_time))
        end_time_str = str(timedelta(seconds=end_time))
        timestamp = f"{start_time_str} - {end_time_str}"
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(timestamp)
        run.bold = True
      doc.add_paragraph(segment['text'])

    # Save the Word document
    download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    sanitized_title = re.sub(r'[<>:"/\\|?*]', '', video_id)
    word_path = os.path.join(download_dir, f"{sanitized_title}.docx")
    doc.save(word_path)

    return word_path

  except Exception as e:
    st.error(f"Error getting text: {str(e)}")
    return None

def main():
    st.title("YouTube Video Downloader and Transcript Generator")

    video_url = st.text_input("Paste YouTube video URL:")
    if video_url:
        video_id = extract_video_id(video_url)
        if video_id:
            st.video(f"https://www.youtube.com/embed/{video_id}")

    st.sidebar.title("Select any option to proceed:")
    video_downloaded = False
    if st.sidebar.button("Download Video"):
        if not video_url:
            st.warning("Please enter a YouTube URL to download the video.")
        else:
            selected_quality = st.selectbox("Select Video Quality:", ["720p", "1080p", "480p", "360p", "240p", "144p"])
            if selected_quality:
                with st.spinner("Downloading Video..."):
                    downloaded_file_path = download_youtube_video(video_url, selected_quality)
                    if downloaded_file_path:
                        video_downloaded = True
    if video_downloaded:
        # Provide a download button for the user
        file_content = open(downloaded_file_path, 'rb').read()
        st.download_button(
            label="Click here to download",
            data=file_content,
            file_name=os.path.basename(downloaded_file_path),
            key="download_button",
        )
    audio_downloaded = False
    if st.sidebar.button("Download Audio"):
        if not video_url:
            st.warning("Please enter a YouTube URL to download the Audio.")
        else:
            with st.spinner("Downloading Audio..."):
                downloaded_audio_path = download_youtube_audio(video_url)
                if downloaded_audio_path:
                    audio_downloaded = True
    if audio_downloaded:
        # Provide a download button for the user
        file_content = open(downloaded_audio_path, 'rb').read()
        st.download_button(
            label="Click here to download",
            data=file_content,
            file_name=os.path.basename(downloaded_audio_path),
            key="download_button",
        )

    transcript_downloaded = False
    if st.sidebar.button("Generate Transcript PDF"):
        if not video_url:
            st.warning("Please enter a YouTube URL to generate the transcript PDF.")
        else:
            with st.spinner("Generating Transcript PDF..."):
                pdf_file_path = get_text_as_word_doc(video_url)  # Ensure this function returns the file path
                if pdf_file_path:
                    transcript_downloaded = True
                    st.success("Transcript PDF generated successfully.")

    if transcript_downloaded:
        # Provide a download button for the transcript PDF
        with open(pdf_file_path, 'rb') as pdf_file:
            pdf_file_content = pdf_file.read()
            st.download_button(
                label="Click here to download the transcript PDF",
                data=pdf_file_content,
                file_name=os.path.basename(pdf_file_path),
                key="transcript_download_button",
                mime='application/pdf'
            )

if __name__ == "__main__":
    main()
